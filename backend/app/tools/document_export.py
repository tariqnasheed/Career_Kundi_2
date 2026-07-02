"""
tools/document_export.py
=============================
Deterministic CV document export — Markdown, PDF (via weasyprint), and DOCX
(via python-docx) — all built from the SAME `GeneratedCV.rendered_content`
JSON produced by `app.agents.cv_builder.render.render_cv()`.

Per the platform-wide design principle ("deterministic, non-generative
transformations happen OUTSIDE the agent graph"), nothing in this module
calls an LLM, a Guardrail, or a Reflector — it is pure template/formatting
code operating on already-approved content. This is what lets export be
instant and free of any token cost no matter how many times a user
re-exports the same CV in a different format.

All three formats are built from the same Markdown intermediate
representation (`build_markdown()`) except DOCX, which needs native
paragraph/heading objects rather than HTML/Markdown text, so it walks the
same `rendered_content` structure directly.
"""

from __future__ import annotations

import io
from typing import Any

import markdown2

from app.agents.job_search.knowledge.study_sources import render_study_source_markdown

_MONTH_ABBR = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]


def _format_date(value: Any) -> str:
    """Render an ISO `YYYY-MM-DD` date string (as produced by `model_dump(mode='json')`) as 'Mon YYYY'. Falls back to the raw value for anything that doesn't parse — never raises on bad/missing input."""
    if not value:
        return ""
    text = str(value)
    try:
        year, month = text[:4], text[5:7]
        return f"{_MONTH_ABBR[int(month) - 1]} {year}"
    except (ValueError, IndexError):
        return text


def _date_range(start: Any, end: Any, is_current: bool = False) -> str:
    if not start and not end:
        return ""
    start_label = _format_date(start)
    end_label = "Present" if is_current else _format_date(end)
    return f"{start_label} – {end_label}".strip(" –")


def _render_generic_entry_lines(entry: dict) -> list[str]:
    """
    Render one entry of a pass-through section (certifications, publications,
    languages, volunteer, awards, references) by picking whichever fields are
    actually populated, rather than assuming every section shares one schema
    — these five sections each have genuinely different field names (see
    `app/agents/cv_builder/render.py`'s module docstring).
    """
    primary = entry.get("name") or entry.get("title") or entry.get("role") or "Untitled"
    secondary_keys = (
        "issuing_organization",
        "organization",
        "publisher",
        "proficiency",
        "relationship_to_user",
        "email",
        "phone",
    )
    secondary = next((entry.get(k) for k in secondary_keys if entry.get(k)), None)
    date_keys = ("issue_date", "date_received", "publication_date", "start_date")
    date_field = next((entry.get(k) for k in date_keys if entry.get(k)), None)

    header = f"**{primary}**"
    if secondary:
        header += f" — {secondary}"
    if date_field:
        header += f" ({_format_date(date_field)})"

    lines = [header]
    description = entry.get("description") or entry.get("abstract")
    if description:
        lines.append(f"  {description}")
    for bullet in entry.get("description_bullets") or []:
        lines.append(f"  - {bullet}")
    return lines


def _render_titled_entry_lines(entry: dict, section_id: str) -> list[str]:
    """Shared renderer for the two sections the agent pipeline actually rewrites: experience and projects."""
    if section_id == "experience":
        title_line = " — ".join(p for p in (entry.get("job_title"), entry.get("company_name")) if p)
    else:
        title_line = entry.get("title", "")

    lines = [f"**{title_line or 'Untitled'}**"]

    meta_bits = [b for b in (entry.get("location"), entry.get("employment_type"), entry.get("role")) if b]
    date_range = _date_range(entry.get("start_date"), entry.get("end_date"), entry.get("is_current", False))
    if date_range:
        meta_bits.append(date_range)
    if meta_bits:
        lines.append("_" + " · ".join(meta_bits) + "_")

    if section_id == "projects" and entry.get("description"):
        lines.append(entry["description"])
    if entry.get("technologies"):
        lines.append("Technologies: " + ", ".join(entry["technologies"]))
    for bullet in entry.get("bullets") or []:
        lines.append(f"- {bullet}")
    return lines


def _render_section_markdown(section: dict) -> list[str]:
    section_id = section.get("section_id", "")
    out = [f"## {section.get('title', 'Section')}"]

    if section_id == "summary":
        out.append(section.get("content", ""))
        return out

    if section_id == "skills":
        items = section.get("items") or []
        if items:
            out.append(", ".join(items))
        return out

    if section_id in ("experience", "projects"):
        for entry in section.get("entries", []):
            out.extend(_render_titled_entry_lines(entry, section_id))
            out.append("")
        return out

    if section_id == "education":
        for entry in section.get("entries", []):
            header = " — ".join(p for p in (entry.get("degree"), entry.get("field_of_study")) if p)
            out.append(f"**{header or entry.get('institution', 'Untitled')}**")
            if entry.get("institution") and header:
                out.append(entry["institution"])
            date_range = _date_range(entry.get("start_date"), entry.get("end_date"), entry.get("is_current", False))
            if date_range:
                out.append(f"_{date_range}_")
            if entry.get("grade"):
                out.append(f"Grade: {entry['grade']}")
            for bullet in entry.get("description_bullets") or []:
                out.append(f"- {bullet}")
            out.append("")
        return out

    if section_id.startswith("custom-"):
        if section.get("free_text_content"):
            out.append(section["free_text_content"])
        if section.get("tags"):
            out.append(", ".join(section["tags"]))
        for entry in section.get("entries", []):
            header = f"**{entry.get('entry_title', 'Untitled')}**"
            if entry.get("subtitle"):
                header += f" — {entry['subtitle']}"
            out.append(header)
            for bullet in entry.get("description_bullets") or []:
                out.append(f"- {bullet}")
        return out

    # Generic pass-through sections: certifications, publications, languages, volunteer, awards, references.
    for entry in section.get("entries", []):
        out.extend(_render_generic_entry_lines(entry))
    return out


def build_markdown(rendered_content: dict[str, Any]) -> str:
    """Deterministic Markdown rendering of `GeneratedCV.rendered_content` — the shared intermediate representation behind every export format below."""
    info = rendered_content.get("personal_info", {})
    lines = [f"# {info.get('full_name') or 'Untitled'}"]

    contact_bits = [b for b in (info.get("headline"), info.get("location"), info.get("email"), info.get("phone")) if b]
    if contact_bits:
        lines.append(" | ".join(contact_bits))
    if info.get("links"):
        link_bits = [f"[{l.get('label', 'Link')}]({l['url']})" for l in info["links"] if l.get("url")]
        if link_bits:
            lines.append(" | ".join(link_bits))
    lines.append("")

    for section in rendered_content.get("sections", []):
        lines.extend(_render_section_markdown(section))
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def export_markdown(rendered_content: dict[str, Any]) -> bytes:
    return build_markdown(rendered_content).encode("utf-8")


# --- PDF export (weasyprint) -----------------------------------------------------------

TEMPLATE_STYLES: dict[str, str] = {
    "modern": """
        body { font-family: 'Helvetica Neue', Arial, sans-serif; color: #1f2937; margin: 40px; line-height: 1.5; }
        h1 { font-size: 26px; margin-bottom: 4px; color: #111827; }
        h2 { font-size: 15px; text-transform: uppercase; letter-spacing: 0.05em; color: #2563eb;
             border-bottom: 1px solid #e5e7eb; padding-bottom: 4px; margin-top: 22px; }
        em { color: #6b7280; font-style: normal; font-size: 13px; }
        ul { margin: 4px 0 10px 18px; padding: 0; }
        li { margin-bottom: 3px; }
    """,
    "classic": """
        body { font-family: 'Georgia', 'Times New Roman', serif; color: #1a1a1a; margin: 48px; line-height: 1.6; }
        h1 { font-size: 24px; text-align: center; margin-bottom: 2px; }
        h2 { font-size: 14px; text-transform: uppercase; border-bottom: 2px solid #1a1a1a;
             padding-bottom: 2px; margin-top: 20px; }
        em { color: #444; font-style: italic; }
        ul { margin: 4px 0 10px 20px; }
    """,
    "compact": """
        body { font-family: Arial, sans-serif; color: #111; margin: 24px; line-height: 1.3; font-size: 12px; }
        h1 { font-size: 20px; margin-bottom: 2px; }
        h2 { font-size: 12px; text-transform: uppercase; color: #374151; margin-top: 12px; margin-bottom: 4px; }
        em { color: #555; font-style: normal; font-size: 11px; }
        ul { margin: 2px 0 6px 16px; }
        li { margin-bottom: 1px; }
    """,
    "creative": """
        body { font-family: 'Trebuchet MS', sans-serif; color: #1f2937; margin: 40px; line-height: 1.5; }
        h1 { font-size: 30px; color: #7c3aed; margin-bottom: 4px; }
        h2 { font-size: 15px; color: #ffffff; background: #7c3aed; display: inline-block;
             padding: 2px 10px; border-radius: 4px; margin-top: 20px; }
        em { color: #7c3aed; font-style: normal; font-size: 13px; }
        ul { margin: 4px 0 10px 18px; }
    """,
}


def _build_html(rendered_content: dict[str, Any], template: str) -> str:
    body_html = markdown2.markdown(build_markdown(rendered_content))
    css = TEMPLATE_STYLES.get(template, TEMPLATE_STYLES["modern"])
    return f"<html><head><meta charset='utf-8'><style>{css}</style></head><body>{body_html}</body></html>"


def export_pdf(rendered_content: dict[str, Any], template: str = "modern") -> bytes:
    from weasyprint import HTML  # imported lazily — weasyprint pulls in native cairo/pango bindings

    html = _build_html(rendered_content, template)
    return HTML(string=html).write_pdf()


# --- DOCX export (python-docx) ---------------------------------------------------------


def _render_titled_entry_docx(doc, entry: dict, section_id: str) -> None:
    if section_id == "experience":
        title_line = " — ".join(p for p in (entry.get("job_title"), entry.get("company_name")) if p)
    else:
        title_line = entry.get("title", "")
    paragraph = doc.add_paragraph()
    paragraph.add_run(title_line or "Untitled").bold = True

    meta_bits = [b for b in (entry.get("location"), entry.get("employment_type"), entry.get("role")) if b]
    date_range = _date_range(entry.get("start_date"), entry.get("end_date"), entry.get("is_current", False))
    if date_range:
        meta_bits.append(date_range)
    if meta_bits:
        doc.add_paragraph(" · ".join(meta_bits))

    if section_id == "projects" and entry.get("description"):
        doc.add_paragraph(entry["description"])
    if entry.get("technologies"):
        doc.add_paragraph("Technologies: " + ", ".join(entry["technologies"]))
    for bullet in entry.get("bullets") or []:
        doc.add_paragraph(bullet, style="List Bullet")


def _render_section_docx(doc, section: dict) -> None:
    section_id = section.get("section_id", "")
    doc.add_heading(section.get("title", "Section"), level=1)

    if section_id == "summary":
        doc.add_paragraph(section.get("content", ""))
        return

    if section_id == "skills":
        items = section.get("items") or []
        if items:
            doc.add_paragraph(", ".join(items))
        return

    if section_id in ("experience", "projects"):
        for entry in section.get("entries", []):
            _render_titled_entry_docx(doc, entry, section_id)
        return

    if section_id == "education":
        for entry in section.get("entries", []):
            header = " — ".join(p for p in (entry.get("degree"), entry.get("field_of_study")) if p)
            paragraph = doc.add_paragraph()
            paragraph.add_run(header or entry.get("institution", "Untitled")).bold = True
            if entry.get("institution") and header:
                doc.add_paragraph(entry["institution"])
            date_range = _date_range(entry.get("start_date"), entry.get("end_date"), entry.get("is_current", False))
            if date_range:
                doc.add_paragraph(date_range)
            if entry.get("grade"):
                doc.add_paragraph(f"Grade: {entry['grade']}")
            for bullet in entry.get("description_bullets") or []:
                doc.add_paragraph(bullet, style="List Bullet")
        return

    if section_id.startswith("custom-"):
        if section.get("free_text_content"):
            doc.add_paragraph(section["free_text_content"])
        if section.get("tags"):
            doc.add_paragraph(", ".join(section["tags"]))
        for entry in section.get("entries", []):
            paragraph = doc.add_paragraph()
            paragraph.add_run(entry.get("entry_title", "Untitled")).bold = True
            if entry.get("subtitle"):
                paragraph.add_run(f" — {entry['subtitle']}")
            for bullet in entry.get("description_bullets") or []:
                doc.add_paragraph(bullet, style="List Bullet")
        return

    # Generic pass-through sections: certifications, publications, languages, volunteer, awards, references.
    for entry in section.get("entries", []):
        for line in _render_generic_entry_lines(entry):
            text = line.strip().lstrip("-").strip()
            if not text:
                continue
            if line.strip().startswith("- "):
                doc.add_paragraph(text, style="List Bullet")
            else:
                doc.add_paragraph(text.replace("**", ""))


def export_docx(rendered_content: dict[str, Any]) -> bytes:
    from docx import Document  # imported lazily, same rationale as weasyprint above

    doc = Document()
    info = rendered_content.get("personal_info", {})
    doc.add_heading(info.get("full_name") or "Untitled", level=0)

    contact_bits = [b for b in (info.get("headline"), info.get("location"), info.get("email"), info.get("phone")) if b]
    if contact_bits:
        doc.add_paragraph(" | ".join(contact_bits))
    if info.get("links"):
        link_bits = [f"{l.get('label', 'Link')}: {l['url']}" for l in info["links"] if l.get("url")]
        if link_bits:
            doc.add_paragraph(" | ".join(link_bits))

    for section in rendered_content.get("sections", []):
        _render_section_docx(doc, section)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# --- Interview pack export (PDF) -------------------------------------------------------

_INTERVIEW_PACK_CSS = """
    body { font-family: 'Helvetica Neue', Arial, sans-serif; color: #1f2937; margin: 36px; line-height: 1.55; font-size: 11pt; }
    h1 { font-size: 22px; color: #5b21b6; margin-bottom: 4px; }
    h2 { font-size: 14px; color: #ffffff; background: #7c3aed; display: inline-block; padding: 3px 10px; border-radius: 4px; margin-top: 22px; }
    h3 { font-size: 13px; color: #374151; border-bottom: 1px solid #e5e7eb; padding-bottom: 3px; margin-top: 16px; }
    h4 { font-size: 12px; color: #6d28d9; margin-top: 12px; margin-bottom: 4px; }
    .meta { color: #6b7280; font-size: 10pt; margin-bottom: 16px; }
    .study-box { background: #f5f3ff; border: 1px solid #ddd6fe; border-radius: 8px; padding: 12px 14px; margin: 10px 0; }
    ul { margin: 4px 0 10px 18px; }
    li { margin-bottom: 4px; }
    .answer { background: #fafafa; border-left: 3px solid #8b5cf6; padding: 10px 12px; margin: 8px 0; }
"""


def build_interview_pack_markdown(
    *,
    job_title: str,
    company_name: str | None,
    questions: list[dict],
    generated_at=None,
    confidence_score: float | None = None,
    role_overview: dict | None = None,
) -> str:
    """Deterministic Markdown for a full interview pack including study material per question."""
    lines = [f"# Interview Pack — {job_title}"]
    if company_name:
        lines.append(f"**Company:** {company_name}")
    meta_bits = []
    if generated_at:
        meta_bits.append(f"Generated: {generated_at}")
    if confidence_score is not None:
        meta_bits.append(f"Confidence: {int(confidence_score * 100)}%")
    if meta_bits:
        lines.append(" | ".join(str(b) for b in meta_bits))
    lines.append("")
    lines.append("> Comprehensive Q&A with zero-prior-knowledge study material for each question.")
    lines.append("")

    if role_overview:
        lines.append("## Role overview")
        if role_overview.get("summary"):
            lines.append(role_overview["summary"])
        if role_overview.get("responsibilities"):
            lines.append("**Key responsibilities**")
            for r in role_overview["responsibilities"][:12]:
                lines.append(f"- {r}")
        if role_overview.get("required_skills"):
            lines.append("**Required skills:** " + ", ".join(role_overview["required_skills"][:20]))
        if role_overview.get("what_employers_expect"):
            lines.append("")
            lines.append("## Employer expectations")
            for item in role_overview["what_employers_expect"][:12]:
                lines.append(f"- {item}")
        if role_overview.get("skill_clusters"):
            lines.append("")
            lines.append("## Skill map")
            for cluster in role_overview["skill_clusters"][:20]:
                lines.append(f"- {cluster}")
        lines.append("")

    for i, q in enumerate(questions, start=1):
        if q.get("export_blocked"):
            continue
        qid = q.get("question_id") or f"Q{i:03d}"
        skill = q.get("skill_tag") or q.get("category", "General")
        lines.append(f"## {qid}: {q.get('question', '')}")
        lines.append(f"**Category:** {q.get('category', 'n/a')} · **Skill:** {skill} · **Difficulty:** {q.get('difficulty', 'Medium')}")
        if q.get("related_skills"):
            lines.append("**Related skills:** " + ", ".join(q["related_skills"]))
        lines.append("")
        lines.extend(_study_material_md_sections(q.get("study_material") or {}))
        lines.extend(render_study_source_markdown(q.get("study_sources")))
        lines.append("### Model answer")
        lines.append(q.get("model_answer") or "\n".join(f"- {p}" for p in (q.get("ideal_answer_points") or [])))
        if q.get("answer_explanation"):
            lines.append("")
            lines.append("### Answer explanation")
            lines.append(q["answer_explanation"])
        lines.append("")
        criteria = q.get("evaluation_criteria") or q.get("ideal_answer_points") or []
        if criteria:
            lines.append("**What interviewers look for**")
            for c in criteria:
                lines.append(f"- {c}")
        mistakes = q.get("common_mistakes") or []
        if mistakes:
            lines.append("**Common mistakes**")
            for m in mistakes:
                lines.append(f"- {m}")
        follow_ups = q.get("follow_up_questions") or q.get("follow_ups") or []
        if follow_ups:
            lines.append("**Follow-up questions**")
            for f in follow_ups:
                lines.append(f"- {f}")
        if q.get("practice_tasks"):
            lines.append("**Practice tasks**")
            for p in q["practice_tasks"]:
                lines.append(f"- {p}")
        lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def _study_material_md_sections(study: dict) -> list[str]:
    """Render structured study material blocks for markdown/PDF export."""
    if not study:
        return []
    lines: list[str] = []
    lines.append("### Study material")
    lines.append("")

    core_parts: list[str] = []
    if study.get("what_this_question_tests"):
        core_parts.append(study["what_this_question_tests"])
    if study.get("overview"):
        core_parts.append(study["overview"])
    if study.get("beginner_explanation"):
        core_parts.append(study["beginner_explanation"])
    if study.get("intermediate_explanation"):
        core_parts.append(study["intermediate_explanation"])
    if study.get("definitions"):
        for definition in study["definitions"]:
            core_parts.append(
                f"**{definition.get('term', 'Term')}** means {definition.get('definition', '')}"
            )
    if study.get("key_concepts"):
        core_parts.append("Key concepts: " + ", ".join(study["key_concepts"]))
    if study.get("explanations"):
        core_parts.extend(str(item) for item in study["explanations"])
    if study.get("principles"):
        core_parts.extend(f"Principle: {item}" for item in study["principles"])
    if core_parts:
        lines.append("**Core idea:**")
        lines.extend(core_parts)
        lines.append("")

    apply_lines: list[str] = []
    if study.get("step_by_step_method"):
        for index, step in enumerate(study["step_by_step_method"], 1):
            apply_lines.append(f"{index}. {step}")
    elif study.get("step_by_step_breakdown"):
        for index, step in enumerate(study["step_by_step_breakdown"], 1):
            apply_lines.append(f"{index}. {step}")
    if study.get("practical_example"):
        apply_lines.append(study["practical_example"])
    if study.get("worked_example"):
        apply_lines.append(study["worked_example"])
    if study.get("skill_explanations"):
        for item in study["skill_explanations"]:
            apply_lines.append(f"{item.get('skill', 'Skill')}: {item.get('explanation', '')}")
    if study.get("what_you_need_to_know_first"):
        for item in study["what_you_need_to_know_first"]:
            apply_lines.append(f"- {item}")
    if study.get("key_terms"):
        apply_lines.append("Key checks: " + ", ".join(str(term) for term in study["key_terms"][:8]))
    if apply_lines:
        lines.append("**How to apply it:**")
        lines.extend(apply_lines)
        lines.append("")

    mistakes = study.get("common_mistakes") or []
    if mistakes:
        lines.append("**Common mistakes:**")
        for mistake in mistakes:
            lines.append(f"- {mistake}")
        lines.append("")

    tips: list[str] = []
    for key in ("how_to_answer_better", "interview_traps", "mini_practice_task"):
        value = study.get(key)
        if isinstance(value, list):
            tips.extend(str(item) for item in value)
        elif value:
            tips.append(str(value))
    if study.get("practice_exercises"):
        tips.extend(f"Practice: {item}" for item in study["practice_exercises"])
    if study.get("revision_notes"):
        tips.extend(str(item) for item in study["revision_notes"])
    if tips:
        lines.append("**Interview tip:**")
        for tip in tips[:5]:
            lines.append(f"- {tip}")
        lines.append("")

    standards = study.get("formula_or_framework") or []
    if isinstance(standards, str):
        standards = [standards]
    checklist = study.get("troubleshooting_checklist") or []
    compliance_notes = [str(item) for item in standards if item]
    compliance_notes.extend(
        str(item)
        for item in checklist
        if item
        and any(
            token in str(item).lower()
            for token in ("standard", "regulation", "compliance", "safety", "haccp", "bs ", "sop")
        )
    )
    if compliance_notes:
        lines.append("**Standards / safety / compliance note:**")
        for note in compliance_notes[:5]:
            lines.append(f"- {note}")
        lines.append("")

    if study.get("related_concepts"):
        lines.append("**Related concepts to study next:** " + ", ".join(study["related_concepts"]))
        lines.append("")

    doc_support = study.get("document_library_support") or {}
    if doc_support:
        lines.append("### Document-library support")
        lines.append("")
        summary = doc_support.get("summary") or (
            "Saved project material matched this question through role/skill overlap."
        )
        lines.append(summary)
        lines.append("")
        if doc_support.get("source_path"):
            lines.append(f"- Source: `{doc_support['source_path']}`")
        if doc_support.get("matched_skills"):
            lines.append("- Matched skills: " + ", ".join(str(s) for s in doc_support["matched_skills"]))
        if doc_support.get("supporting_focus"):
            lines.append(
                "- Supporting focus: " + ", ".join(str(s) for s in doc_support["supporting_focus"])
            )
        for snippet in doc_support.get("snippets") or []:
            if snippet and snippet.strip():
                lines.append(f"- Snippet: {snippet}")
        lines.append("")

    return lines


def build_study_material_markdown(
    *,
    job_title: str,
    questions: list[dict],
    generated_at=None,
    role_overview: dict | None = None,
) -> str:
    """Study-material-only PDF — one deep module per question."""
    lines = [f"# Study Material — {job_title}", ""]
    if generated_at:
        lines.append(f"*Generated: {generated_at}*")
    lines.append("")
    if role_overview and role_overview.get("summary"):
        lines.append(role_overview["summary"])
        lines.append("")
    for i, q in enumerate(questions, start=1):
        if q.get("export_blocked"):
            continue
        qid = q.get("question_id") or f"Q{i:03d}"
        lines.append(f"## {qid}: Study guide for this question")
        lines.append(f"*Question:* {q.get('question', '')}")
        lines.append("")
        lines.extend(_study_material_md_sections(q.get("study_material") or {}))
        lines.extend(render_study_source_markdown(q.get("study_sources")))
        lines.append("---")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def build_questions_answers_markdown(
    *,
    job_title: str,
    company_name: str | None,
    questions: list[dict],
    generated_at=None,
) -> str:
    """Q&A-only PDF without study sections."""
    lines = [f"# Questions & Answers — {job_title}"]
    if company_name:
        lines.append(f"**Company:** {company_name}")
    if generated_at:
        lines.append(f"*Generated: {generated_at}*")
    lines.append("")
    for i, q in enumerate(questions, start=1):
        if q.get("export_blocked"):
            continue
        qid = q.get("question_id") or f"Q{i:03d}"
        lines.append(f"## {qid}")
        lines.append(f"**Q:** {q.get('question', '')}")
        lines.append("")
        lines.append("**Model answer:**")
        lines.append(q.get("model_answer") or "")
        if q.get("answer_explanation"):
            lines.append("")
            lines.append("**Explanation:**")
            lines.append(q["answer_explanation"])
        lines.append("")
        lines.append("---")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def _render_pdf(md: str, css: str = _INTERVIEW_PACK_CSS) -> bytes:
    from weasyprint import HTML

    body_html = markdown2.markdown(md, extras=["fenced-code-blocks", "tables"])
    html = f"<html><head><meta charset='utf-8'><style>{css}</style></head><body>{body_html}</body></html>"
    return HTML(string=html).write_pdf()


def export_interview_pack_pdf(
    *,
    job_title: str,
    company_name: str | None,
    questions: list[dict],
    generated_at=None,
    confidence_score: float | None = None,
    role_overview: dict | None = None,
) -> bytes:
    md = build_interview_pack_markdown(
        job_title=job_title,
        company_name=company_name,
        questions=questions,
        generated_at=generated_at,
        confidence_score=confidence_score,
        role_overview=role_overview,
    )
    return _render_pdf(md)


def export_study_material_pdf(
    *,
    job_title: str,
    company_name: str | None = None,
    questions: list[dict],
    generated_at=None,
    confidence_score: float | None = None,
    role_overview: dict | None = None,
) -> bytes:
    del company_name, confidence_score  # study PDF is role-focused
    md = build_study_material_markdown(
        job_title=job_title,
        questions=questions,
        generated_at=generated_at,
        role_overview=role_overview,
    )
    return _render_pdf(md)


def export_questions_answers_pdf(
    *,
    job_title: str,
    company_name: str | None,
    questions: list[dict],
    generated_at=None,
    confidence_score: float | None = None,
    role_overview: dict | None = None,
) -> bytes:
    del confidence_score, role_overview
    md = build_questions_answers_markdown(
        job_title=job_title,
        company_name=company_name,
        questions=questions,
        generated_at=generated_at,
    )
    return _render_pdf(md)
